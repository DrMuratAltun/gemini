
import streamlit as st
import google.generativeai as genai
import markdown
from docx import Document
from bs4 import BeautifulSoup
from PyPDF2 import PdfFileReader
import tempfile
import os

# Configure the API key
genai.configure(api_key=os.getenv('gemini_api'))

# Function to convert PDF to text
def pdf_to_text(file):
    with open(file, 'rb') as f:
        pdf = PdfFileReader(f)
        text = ""
        for page_num in range(pdf.numPages):
            page = pdf.getPage(page_num)
            text += page.extract_text()
    return text

# Function to upload file to the Generative AI API
def upload_file(file_path):
    st.write("Uploading file...")
    text_file = genai.upload_file(path=file_path)
    st.write(f"Completed upload: {text_file.uri}")
    return text_file

# Function to convert text to Markdown
def to_markdown(text):
    text = text.replace('•', '  *')
    return textwrap.indent(text, '> ', predicate=lambda _: True)

chat_session = None

# Function to build the model
def build_model(text_file):
    global chat_session
    generation_config = {
        "temperature": 0.2,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }

    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        generation_config=generation_config,
        system_instruction="""Yüklenen belgedeki bilgilere göre Türkçe cevap ver.
        Eğer sorunun cevabı belgede bulunmuyorsa 'Belgede Cevap Bulunmuyor' yaz.
        """,
    )

    chat_session = model.start_chat(history=[])
    response = chat_session.send_message(["Yüklenen belgeyi bir cümle ile özetle", text_file])

    st.markdown(to_markdown(response.text))

# Function to interact with the chat model
def chat(prompt):
    try:
        response = chat_session.send_message(prompt)
        markdown_text = to_markdown(response.text)
        st.markdown(markdown_text)
        return response.text
    except ValueError:
        st.write(response.prompt_feedback)
        st.write(response.candidates[0].finish_reason)
        st.write(response.candidates[0].safety_ratings)
    except Exception as e:
        st.write("An unexpected error occurred:", e)

# Function to generate a report based on questions
def generate_report(questions):
    report_text = "\n## SORULARINIZ VE CEVAPLARI\n"
    for question in questions:
        report_text += f"\n## {question}\n"
        answer = chat(question)
        report_text += f"\n{answer}\n"
    return report_text

# Function to convert Markdown to HTML
def convert_Markdown_to_HTML(report_text):
    html_text = markdown.markdown(report_text)
    return html_text

# Function to add HTML to a Word document
def add_html_to_word(html_text, doc):
    soup = BeautifulSoup(html_text, 'html.parser')
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            doc.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            doc.add_heading(element.get_text(), level=6)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name:
            doc.add_paragraph(element.get_text())

# Streamlit interface
st.title("REPORT GENERATOR: ASK YOUR QUESTIONS TO A PDF FILE BY MURAT KARAKAYA AKADEMI")
st.write("Upload a PDF to ask questions and get the answers.")

uploaded_file = st.file_uploader("Upload PDF", type="pdf")
questions_input = st.text_area("Enter Questions", placeholder="Type your questions here, one per line.", height=150)

if uploaded_file and questions_input:
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(uploaded_file.read())
        temp_file_path = temp_file.name

    text_content = pdf_to_text(temp_file_path)
    text_file = upload_file(temp_file_path)
    build_model(text_file)

    questions = questions_input.split("\n")
    report_text = generate_report(questions)

    html_text = convert_Markdown_to_HTML(report_text)
    doc = Document()
    add_html_to_word(html_text, doc)
    
    doc_name = os.path.basename(temp_file_path).replace(".pdf", ".docx")
    doc_name = "Rapor " + doc_name
    doc.save(doc_name)
    
    st.markdown(report_text)
    st.write("Document generated successfully!")
    
    with open(doc_name, "rb") as file:
        st.download_button(label="Download Report", data=file, file_name=doc_name)
    
    os.remove(temp_file_path)
    os.remove(doc_name)

    genai.delete_file(text_file.name)